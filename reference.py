

































Request  urls

def your_view(request):
referrer_url = request.META.get('HTTP_REFERER', None)
r request_url = request.build_absolute_uri()




















def get(self, request):
dummy_content_type, _ = ContentType.objects.get_or_create(
app_label='custom_permissions', # Use your app label
model='dummycontenttype', # Use a unique model name
)
view_accounting_permission, _ = Permission.objects.get_or_create(
codename='view_accounting',
name='Can View Accounting',
content_type=dummy_content_type,
)
permission = Permission.objects.get(codename='view_accounting')
context = {
"message": "Done"
}
return Response(permission)




class Accountingpermission(viewsets.ReadOnlyModelViewSet):
serializer_class = AccountingPermissionSerializer
def get_queryset(self):
permission = Permission.objects.get(codename='view_accounting')
serializer = AccountingPermissionSerializer(permission)
return Response(serializer.data)





























from rest_framework.permissions import BasePermission


class UserPermissionlistAPI(viewsets.ReadOnlyModelViewSet):
http_method_names = ["get"]
serializer_class = UserPermissionlistserializer
queryset = Permission.objects.all()


def list(self, request, *args, **kwargs):
queryset = self.filter_queryset(self.get_queryset())
serializer = UserPermissionlistserializer(queryset, many=True)
data = serializer.data


user_permissions = set(request.user.get_all_permissions())


grouped_data = [] # Initialize an empty list to store grouped data


for item in data:
model_name = item['model_name']
# Check if the model_name is already present in the grouped_data list
model_entry = next((entry for entry in grouped_data if entry['model_name'] == model_name), None)


if model_entry is None:
# If the model_name is not present, add it along with the permission details
grouped_data.append({
'model_name': model_name,
'model_permission': [{
'id': item['id'],
'name': item['name'],
'has_permission': self.check_permission(user_permissions, item)
}]
})
else:
# If the model_name is already present, append the permission details to the existing entry
model_entry['model_permission'].append({
'id': item['id'],
'name': item['name'],
'has_permission': self.check_permission(user_permissions, item)
})


# Return the list of dictionaries in the response
return Response(grouped_data)


def check_permission(self, user_permissions, permission_data):
# Check if the user has the specific permission
return f"{permission_data['content_type']}.{permission_data['codename']}" in user_permissions




'has_permission': self.check_permission(user_permissions, item)

{% for course in courses|slice:":5" %}
<li class="f-list"><a href="#courses" class="a_tag">{{ course.name }}</a></li>
{% endfor %}



<li><a href="{{page.get_parent.url}}?category={{ cat.name }}"><span class="blog-tag">{{ cat.name }}</span></a></li>






























Base 64



import base64
import json


# Given payload
payload = {
"org_id": 11,
"startdate": "1111",
"enddate": 2222,
"phonenumber": "987923342442479",
}


# Convert the payload to a JSON string
payload_json = json.dumps(payload)


# Encode the payload using Base64
encoded_payload = base64.b64encode(payload_json.encode())


# Your generated token
token = encoded_payload


print("Generated Token:", token)


key = "98098080"
if key== "98098084340":
decoded_payload = base64.b64decode(token)


# Parse the JSON payload
payload = json.loads(decoded_payload)
# Extract the phone number from the payload
phone_number = payload["phonenumber"]


print("Phone Number-------------------------------------------------------------:", phone_number)
else:
print("----------------------invlid key",)


















Response modify


































































































































































































































































































































































































































































































































































































<!-- <div class="col-12 ">
<div clas="heading">
<h3 class="text-dark"><i>Response :</i> </h3>
</div>


<table class="table">
<thead>
<tr>
<th scope="col">Question</th>
<th scope="col">Answer</th>
</tr>
</thead>
{% if summary %}
<tbody>
<tr>
<td>
<div>
<p>Summary: {{ response.summary_type}}</p>
<p>Instructions: {{ response.summary_instruction }}</p>
</div>
</td>
<td> {% if response.summary %} {{ response.summary }} {% else %} N/A {% endif %}</td>
</tr>
</tbody>
{% endif %}
{% if question %}
<tbody>
<tr>
<td>
<div>
<p>Question: {{ response.question}}</p>
<p>Quesion_instruction: {{response.quesion_instruction}}</p>
<p>Question_keyword: {{response.question_keyword}}</p>
</div>
</td>
<td> {% if response.answer %} {{ response.answer }} {% else %} N/A {% endif %}</td>
</tr>
</tbody>
{% endif %}
{% if theme %}
<tbody>
<tr>
<td>
<div>
<p>Theme_type: {{ result.theme_type}}</p>
<p>Instruction: {{result.instruction}}</p>
</div>
</td>
<td> {% if response %} {{ response|safe }} {% else %} N/A {% endif %}</td>
</tr>
</tbody>
{% endif %}
{% if frequency %}
<tbody>
<tr>
<td>
<div>
<p>Instruction: {{ response.instruction}}</p>
</div>
</td>
<td> {% if response.frequency %} {{ response.frequency }} {% else %} N/A {% endif %}</td>
</tr>
</tbody>
{% endif %}
{% if compareviewpoint %}
<tbody>
<tr>
<td>
<div>
<p>Question: {{ response.question}}</p>
<p>Instruction: {{ response.instruction}}</p>
<p>keywords: {{ response.keywords}}</p>


</div>
</td>
<td> {% if response.compare_viewpoint %} {{ response.compare_viewpoint }} {% else %} N/A {% endif %}</td>
</tr>
</tbody>
{% endif %}
{% if excel_column_themantic %}
<tbody>
<tr>
<td>
<div>
<p>Question: {{ response.question}}</p>
<p>Instruction: {{ response.theme_type}}</p>
<p>keywords: {{ response.theme_instruction}}</p>
</div>
</td>
<td> {% if response.answer %} {{ response.answer }} {% else %} N/A {% endif %}</td>
</tr>
</tbody>


{% endif %}
{% if categorize_text_excel %}
<tbody>
<tr>
<td>
<div>
<p>Question: {{ response.question}}</p>
<p>Instruction: {{ response.categorize_instructions}}</p>
</div>
</td>
<td> {% if response.answer %} {{ response.answer }} {% else %} N/A {% endif %}</td>
</tr>
</tbody>
{% endif %}
</table>
</div> -->































def post(self, request):
choice = request.session.get("choice")
if not choice:
errormessage="Select at least one option."
return render(request,'choiceerror.html',{'errormessage':errormessage})
# if choice:
# del request.session["choice"]
if request.user.is_authenticated:
uploaded_file = request.FILES.getlist("file")
if "upload_option" in request.session:
upload_option = request.session["upload_option"]
# del request.session["upload_option"]
else:
upload_option='new'
if upload_option == "previous":
files = request.POST.getlist("files")
uploaded_file = Files.objects.filter(id__in=files)
uploaded_file = [file.file for file in uploaded_file]
else:
uploaded_file = request.FILES.getlist("demo_file")


if not isinstance(uploaded_file, list):
uploaded_file = [uploaded_file]
instance = FileHandler(uploaded_file)
access=self.request.user.is_authenticated
if access:
df = instance.upload_documents(max_documents=50, max_words=100000)
if not access:
df = instance.upload_documents(max_documents=3, max_words=15000)


if choice == Anaylsis.Summarize.value:
forms = SummerizeType(request.POST)
if forms.is_valid():
summary_type = forms.cleaned_data["summary"]
summary_instruction = forms.cleaned_data["instruction"]
indiviual_summary = forms.cleaned_data["individual_summary"]
instance = SumarrizeClass(
is_demo=False,
summary_type=summary_type,
summary_instructions=summary_instruction,
individual_summaries=indiviual_summary,
)
print("summary_type",summary_type)


response,error = instance.call_summarize(df)
if summary_type == "Bullet points":
html_text = markdown.markdown(response, extensions=['extra'])
data = html_text
print("data (Essay):", data)
else:
sections = response.split('\n\n')
formatted_dict = {}
for section in sections:
section_parts = section.split(':', 1)
key = section_parts[0].strip()
value = section_parts[1].strip() if len(section_parts) > 1 else ""
formatted_dict[key] = value
data = formatted_dict
print("data (Formatted Dict):", data)
# if summary_type == "Essay":
# html_text = markdown.markdown(response, extensions=['extra'])


# sections = response.split('\n\n')
# formatted_dict = {}


# for section in sections:
# section_parts = section.split(':', 1)
# key = section_parts[0].strip()
# value = section_parts[1].strip() if len(section_parts) > 1 else ""
# formatted_dict[key] = value




if error:
return render(request, "response.html", {"response": response,'error':error})
if request.user.is_authenticated:
UserQuery.objects.create(
user=self.request.user,
question={
"Summary_type": summary_type,
"Summary_instruction": summary_instruction,
},
answer=response,
)
return render(
request, "response.html", {"response": response}
)
else:
UserQuery.objects.create(
question={
"Summary_type": summary_type,
"Summary_instruction": summary_instruction,
},
answer=response,
)
context = {
"summary_type": summary_type,
"summary_instruction": summary_instruction,
"summary": response,
}
# return JsonResponse({"response":context,'summary':True})
return render(
request,"techdemo.html", {"response": data, "summary": True}
)








































Not found page



CSV READ


Csv read


import csv
from django.core.management.base import BaseCommand
from apps.users.models import UserQuery
from apps.users.models import User


class Command(BaseCommand):
help = "Add email to user from a CSV file"


def handle(self, *args, **options):
csv_file_path = "apps/users/management/commands/AILYZE_user_past_questions_and_answers.csv" # Update the path to your CSV file


# Open the CSV file
with open(csv_file_path, 'r') as file:
csv_reader = csv.reader(file)
header_row = next(csv_reader)
for row in csv_reader:
question = row[1]
answer = row[2]
email = row[3]
user = User.objects.filter(email=email).first()
# # Create a new UserQuery object and save it to the database
user_query,created = UserQuery.objects.get_or_create(question=question, answer=answer,user=user.id)
print(f"New UserQuery created:")










class SalesOrderAPI(APIView):
def get(seif,request,id=None,format=None):
org = request.headers.get("organization")
try:
if settings.DEBUG:
data = {
"count": "0",
"results": []
}
return Response(data, status=status.HTTP_200_OK)
if id is not None:
data = seterasosdetails(id)
count = len(data)
data_obj = {
"count": count,
"results": [{data}]
}
return Response(data_obj, status=status.HTTP_200_OK)
else:
salesorders = seteraSOS(org)
count = len(salesorders)
data = {
"count": count,
"results": [{salesorders}]
}
return Response(data, status=status.HTTP_200_OK)
except Exception as e:
return Response(status=status.HTTP_400_BAD_REQUEST)





Send static url false


DEBUG = False




from django.views.static import serve


 re_path(r'^media/(?P<path>.*)$', serve,{'document_root': settings.MEDIA_ROOT}),
   re_path(r'^static/(?P<path>.*)$', serve,{'document_root': settings.STATIC_ROOT}),
]




STATIC_URL = '/static/'
# # STATICFILES_DIRS = [
# #     BASE_DIR /  "static",
# #     # BASE_DIR / "accounts" / "static"
# # ]
STATIC_ROOT = BASE_DIR / 'static'
























dict Compare

class ProcessQuery(View):
   a = {
       Anaylsis.Summarize.value: lambda request: render(request,'chioceform.html',{'forms':SummerizeType()}),
       Anaylsis.Ask_a_specific_question.value:  lambda request:render(request,'chioceform.html',{'forms':SPecificQuestion()}),
       Anaylsis.Conduct_thematic_analysis.value: lambda request: render(request,'chioceform.html',{'forms':ThemeType()}),
       Anaylsis.Identidy_which_document_contain_a_certain_viewpoint.value:  lambda request: render(request,'chioceform.html',{'forms':IdentifyViewpoint()}),
       Anaylsis.Compare_viewpoints_across_documents.value:  lambda request: render(request,'chioceform.html',{'forms':CompareViewpoint()})
   }


   def post(self,request):
       choice=request.session.get('choice')
       variables = {
       "Summarize": ["summary", "instruction"],
       "Ask_a_specific_question": ["question", "instruction", "keywords"],
       "Conduct_thematic_analysis": ["theme_type", "instruction"],
       "Identidy_which_document_contain_a_certain_viewpoint": ["instruction"],
       "Compare_viewpoints_across_documents": ["instruction", "question"]
   }


       if choice in variables:
           print("-------------------------------------", choice)
           values = [request.POST.get(var, '') for var in variables[choice]]
           print("This is ins=", *values)


      




       render_fun=self.a.get(choice)
       return HttpResponse("Done ")
  










Data enter without form 




def my_view(request):
   if request.method == 'POST':
       # Process the POST data
       data = request.POST.get('input_data')
       # Perform some operations with the data
       result = data.upper()
       return HttpResponse(result)
   else:
       # Handle GET requests
       return HttpResponse("This is a GET request.")














 #  run in shell 
from django.test import RequestFactory
factory = RequestFactory()


request = factory.post('/my-view-url/', {'input_data': 'example data'})
# request = factory.post('/my-view-url/', {'input_data': 'John\nDoe\njohn@example.com\n123 Main St\npassword123'})


request = factory.post('/my-view-url/', {'input_data': 'first_name:John\nlast_name:Doe\nemail:john@example.com\naddress:123 Main St\npassword:password123'})






response = my_view(request)
print(response.content)














Foreign key some values



class SubscriptionResponseSerializer(serializers.ModelSerializer):
   organization = OrgSerializer()
   sim = SIMSeriliazer()
   carrier = CarrierSerializer()
   mobile_product = MobileproductSerializer()
   barring_voice = BarringVoiceSerializer()
   barring_mms = BarringMMSSerializer()
   barring_sms = BarringSMSSerializer()
   barring_roaming = BarringRoamingSerializer()
   barring_data = BarringDataSerializer()
   barring_roamingdata = BarringRoamingDataSerializer()


   class Meta:
       model = Subscription
       fields = list_display = (
           "id",
           "organization",
           "number"
       )


       read_only_fields = ("temp_close",)




   def get_organization(self, obj):
       return {
       "id": obj.organization.id,
       "name": obj.organization.name,
   }


   def to_representation(self, instance):
       representation = super().to_representation(instance)
       representation['organization'] = self.get_organization(instance)
       return representation






























Code creaTE
try:




           dataProduct = {
               "dataCode": obj.mobile_product.datacode,
               "threeGDown": obj.mobile_product.dataspeed3GUp.name,
               "threeGUp": obj.mobile_product.dataspeed3GUp.name,
               "fourGDown": obj.mobile_product.dataspeed4GDown.name,
               "fourGUp": obj.mobile_product.dataspeed4GUp.name,
               "priority": obj.mobile_product.priority,
           }
           subscription_payload = {
               "subscription": obj.number,
               "dataProduct": dataProduct,
               "dataRoamingLimit": "",
               "barrings": "",
               "imsi": obj.sim.imsi,
               "iccID": obj.sim.icc,
               "voicemail": obj.voicemail_number,
           }
           subscriptionId = serviceId.objects.filter(subscription=obj.id).first()


           # tempopen payload


           tempopen_payload = {
               "subscription": obj.number,
               # "subscriptionId": subscriptionId.subscriptionId,
               # "serviceId": subscriptionId.serviceId,
           }
           print("----", tempopen_payload)


           # datapackege paylaod
           datapackage_payload = {
               "subscription": obj.number,
               # "subscriptionId" :subscriptionId.subscriptionId,
               "dataProduct": dataProduct,
           }
           # data_roaming_payload


           data_roaming_payload = {
               "subscription": obj.number,
               # "subscriptionId" :subscriptionId.subscriptionId,
               "dataRoamingLimit": obj.barring_data.operator_code,
           }


       except Exception as e:
           print(e)







     Permission 


from rest_framework.permissions import BasePermission, SAFE_METHODS




class WriteByAdminOnlyPermission(BasePermission):
   def has_permission(self, request, view):
       user = request.user
       if request.method in SAFE_METHODS:
           return True
       if request.method in ("POST", "DELETE", "PUT", "PATCH") and (
           user.is_superuser or user.is_staff
       ):
           return True


       return False




class ReadByAdminOnlyPermission(BasePermission):
   def has_permission(self, request, view):
       user = request.user
       if request.method in ("POST", "GET", "DELETE", "PUT", "PATCH") and (
           user.is_superuser or user.is_staff
       ):
           return True
       return False




class IsCustomAdminUser(BasePermission):
   """
   Allows access only to admin users.
   """


   def has_permission(self, request, view):
       return bool(request.user and request.user.user_type.lower() == "admin")


























Base64 dummy





Video compress 





@receiver(pre_save, sender=PostItem)
def compress_media(sender, instance, **kwargs):
       data = instance.data
       # data = getattr(instance, 'data')
       def filetype(data):
           if data.name.lower().endswith(('.jpg', '.jpeg', '.png', '.gif')):
               return True
           if data.name.lower().endswith(('.mp4', '.mov', '.avi', '.wmv')):
               return False
       status = filetype(data)
       image_field = data
       video_field = data
      # Compress image
       if image_field and status == True:
           im = Image.open(image_field)
           if im.mode == 'RGBA':
               im = im.convert('RGB')
           im_io = BytesIO()
           im.save(im_io, 'JPEG', quality=60)    
           compressed_image = File(im_io, name=image_field.file.name)
           image_field.file = compressed_image


       # Compress video
       #         # Compress video
       if video_field and status==False:
           input_data = video_field.file.read()
           with tempfile.NamedTemporaryFile(delete=False) as temp_file:
               temp_file.write(input_data)
               temp_file_path = temp_file.name
           # Compress video using MoviePy
           video = VideoFileClip(temp_file_path)
           compressed_video = video.resize((640, 360))
           compressed_video_file = tempfile.NamedTemporaryFile(delete=False, suffix='.mp4')
           compressed_video.write_videofile(compressed_video_file.name, codec='libx264', audio_codec='aac')
           # Save compressed video to media folder
           with open(compressed_video_file.name, 'rb') as f:
               compressed_video_file = File(f)
               compressed_video_path = os.path.join('', compressed_video_file.name.split('/')[-1])
               default_storage.save(compressed_video_path, compressed_video_file)
           instance.data = compressed_video_path


           os.unlink(temp_file_path)
           os.unlink(compressed_video_file.name)
















          
# @receiver(pre_save, sender=Post)
# def compress_media(sender, instance, **kwargs):
#     # Check if the instance is an image or video model
#     if getattr(instance, 'image') or getattr(instance, 'video'):
#         image_field = getattr(instance, 'image')
#         video_field = getattr(instance, 'video')


#         # Compress image
#         if image_field:
#             im = Image.open(image_field)
#             if im.mode == 'RGBA':
#                 im = im.convert('RGB')
#             im_io = BytesIO()
#             im.save(im_io, 'JPEG', quality=60)    
#             compressed_image = File(im_io, name=image_field.file.name)
#             image_field.file = compressed_image


#         # Compress video
#         if video_field:
#             input_data = video_field.file.read()
#             with tempfile.NamedTemporaryFile(delete=False) as temp_file:
#                 temp_file.write(input_data)
#                 temp_file_path = temp_file.name
#             # Compress video using MoviePy
#             video = VideoFileClip(temp_file_path)
#             compressed_video = video.resize((640, 360))
#             compressed_video_file = tempfile.NamedTemporaryFile(delete=False, suffix='.mp4')
#             compressed_video.write_videofile(compressed_video_file.name, codec='libx264', audio_codec='aac')
#             # Save compressed video to media folder
#             with open(compressed_video_file.name, 'rb') as f:
#                 compressed_video_file = File(f)
#                 compressed_video_path = os.path.join('video', compressed_video_file.name.split('/')[-1])
#                 default_storage.save(compressed_video_path, compressed_video_file)
#             instance.video.name = compressed_video_path


#             os.unlink(temp_file_path)
#             os.unlink(compressed_video_file.name)

















































if video_field:
           input_data = video_field.file.read()
           with tempfile.NamedTemporaryFile(delete=False) as temp_file:
               temp_file.write(input_data)
               temp_file_path = temp_file.name
           # Compress video using MoviePy
           video = VideoFileClip(temp_file_path)
           compressed_video = video.resize((640, 360))
           compressed_video_file = tempfile.NamedTemporaryFile(delete=False, suffix='.mp4')
           compressed_video.write_videofile(compressed_video_file.name, codec='libx264', audio_codec='aac')
           # Save compressed video to media folder
           with open(compressed_video_file.name, 'rb') as f:
               compressed_video_file = File(f)
               compressed_video_path = os.path.join('RoboKidz/thumbs/', compressed_video_file.name.split('/')[-1])
               default_storage.save(compressed_video_path, compressed_video_file)
       # Update instance video field with compressed video path
           instance.video.name = compressed_video_path
           # Delete temporary files
           os.unlink(temp_file_path)
           os.unlink(compressed_video_file.name)
           # input_data = video_field.file.read()
           # with tempfile.NamedTemporaryFile(delete=False) as f:
           #     f.write(input_data)
           #     f.seek(0)
           #     file_path = f.name
           #     print("===========================",file_path)
           #     input_file = file_path
           #     output_file = 'compressed_video.mp4'
           #     target_bitrate = '1000k'
           #     video = VideoFileClip(input_file)
           #     compressed_video = video.resize(width=640)
           #     # Set the target bitrate
              
           #     compressed_video.write_videofile(output_file, bitrate=target_bitrate)
          




#             print("----------------",video_field.path)
# # Set the input and output file names
#             input_file = video_field
#             output_file = 'compressed_video.mp4'
#             target_bitrate = '1000k'
#             clip = VideoFileClip(input_file)
#             clip_resized = clip.resize(width=640)
#             # Set the target bitrate
#             clip_resized.write_videofile(output_file, bitrate=target_bitrate)




       #     input_data = video_field.file.read()
       #     with tempfile.NamedTemporaryFile(delete=False) as temp_file:
       #         temp_file.write(input_data)
       #         temp_file_path = temp_file.name
              
       #     # # Compress video using MoviePy
       #     video = VideoFileClip(temp_file_path)
       #     compressed_video = video.resize((640, 360))
       #     compressed_video.write_videofile(video_field.file.name)
       #     with open(video_field.file.name, 'rb') as f:
       #         f.seek(0)
       #         compressed_video = File(f)
             
       #     video_field.file = compressed_video
       # os.unlink(video_field.file.name)















Django reference code

python manage.py dumpdata auth.group --indent 2 > apps/group.json


Create default data first time with  migrate





sigal.py
from django.db.models.signals import post_migrate
from django.dispatch import receiver
from app.models import School
from django.contrib.auth.models import Group


@receiver(post_migrate)
def create_default_data(sender, **kwargs):
   if sender.name == 'app':
       l = ["superuser","admin","customer","seller"]
       for i in l:
           Group.objects.get_or_create(name= i)



App.py 
from django.apps import AppConfig




class AppConfig(AppConfig):
   default_auto_field = "django.db.models.BigAutoField"
   name = "app"
   def ready(self):
       from app import signal



































class ListManagementAPI(ModelViewSet):
   serializer_class = ListManagementSerializer
   queryset = ListManagement.objects.all()


   def create(self, request,*args, **kwargs,):
       f= request.data['name']
       decoded_file = f.read().decode()
       io_string = io.StringIO(decoded_file)
       reader = csv.reader(io_string)
       next(reader)
       column_values = []
       for row in reader:
           column_values.append(row[0])  # Change index to match desired column
       print(column_values)
      
       serializer = ListManagementSerializer(data=request.data)
       if serializer.is_valid():
           # try:
               f = serializer.validated_data["name"]
               # decoded_file = f.read().decode()
               # io_string = io.StringIO(decoded_file)
               # reader = csv.reader(io_string)
               # next(reader)
          
               # for row in reader:
               #   print(str(row))
           # except Exception as e:
           #     print(e)
         
           # self.perform_create(serializer)
           # return Response(serializer.data, status=201)
       return Response(serializer.errors, status=400)
  






   def perform_create(self, serializer):
       serializer.save(created_by=self.request.user, updated_by=self.request.user)       


















import pytz
TIMEZONE_CHOICES = []
for tz in pytz.all_timezones:
   TIMEZONE_CHOICES.append((tz, tz))
TIMEZONE_CHOICES = tuple(TIMEZONE_CHOICES)


from django.contrib.auth.models import Group


# get_or_create return error due to
new_group = Group.objects.get_or_create(name = 'groupName')
print(type(new_group))       # return tuple





operation_summary="Sign Up API"








IP ADDRESS (localhost)




def loginhistory(sender,instance, **kwargs):
   ## importing socket module
import socket
## getting the hostname by socket.gethostname() method
hostname = socket.gethostname()
## getting the IP address using socket.gethostbyname() method
ip_address = socket.gethostbyname(hostname)
## printing the hostname and ip_address
print(f"Hostname: {hostname}")
print(f"IP Address: {ip_address}")
   d = str(urlopen('http://checkip.dyndns.com/').read())
   l = r.compile(r'Address: (\d+\.\d+\.\d+\.\d+)').search(d).group(1)
   print("This is ip",l)
  
  print("--------------------",socket.gethostbyname(socket.gethostname()))




 def get_client_ip(self,request):
       x_forwarded_for = request.META.get('HTTP_X_FORWARDED_FOR')
       if x_forwarded_for:
           ip = x_forwarded_for.split(',')[0]
       else:
           ip = request.META.get('REMOTE_ADDR')


       print("This is ip---",ip)




Custom JWT


class LoginAPI(APIView):
   @swagger_auto_schema(request_body=LoginSerializer)
   def post(self, request):
       email = request.data.get('email')
       password = request.data.get('password')
       # user = User.objects.filter(email=email,).first()
       user=authenticate(request=request,email=email,password=password)


       if user:
           user = LoginDetailsSerializer(user)
           auth_token = jwt.encode({'email': user.data.get('email'), 'userid': user.data.get('userid'),
                                    'exp': datetime.datetime.timestamp(
                                        (datetime.datetime.now() + datetime.timedelta(days=1, hours=3)))},
                                   settings.SECRET_KEY, 'HS256')
           data = {
               'user': user.data, 'token': auth_token
           }
           return Response({'data': data, 'issuccess': True}, status=status.HTTP_200_OK)
       return Response({"message": "Invalid Credentials"}, status=status.HTTP_401_UNAUTHORIZED)














USER TYPE

user_type = serializers.ChoiceField(choices=UserType.choices())




def create(self, validated_data):
       validated_data["mobile_verified"] = False
       print("---------------------------------------",UserType.B2B.value)
       validated_data["user_type"] = UserType.B2B.value
       classes = validated_data.pop("classes")
       user = User.objects.create_user(**validated_data)
       gen_pass = "".join((random.choice("abcdxyzpqr1234567890") for i in range(8)))
       print(gen_pass)
       user.password = gen_pass
       # send_passwod(user.email, gen_pass)
       user.set_password(user.password)
       user.save()
       self.create_user_video(user, classes)
       return user



foreign key method 




class kitSerializer(serializers.ModelSerializer):
   qr_code = serializers.CharField(source = 'qr code')
   qr_cetegory = serializers.CharField(source = 'category')
    class Meta:
       model = Kit
       fields = ("id", "qr_cetegory","qr_code" )
       read_only_fields = ()


   






=================================


 def to_representation(self, instance):
       response = super().to_representation(instance)
       response['category'] = CategorySerializer(instance.category,context={"request": self.context["request"]}).data
       return response










Pagination (remove)

"DEFAULT_PAGINATION_CLASS": "rest_framework.pagination.PageNumberPagination",
   "PAGE_SIZE": 10,










class CategoryApiView(ModelViewSet):
   permission_classes = [IsCustomAdminUser]
   # pagination_class = None
   queryset = Category.objects.all()
   serializer_class = CategorySerializer
   filterset_fields = ["name"]
   search_fields = ["name"]
















If else convert on Dictionary

class OrgsForUserView(generics.ListAPIView):
   permission_classes = [IsAuthenticated]
   authentication_classes = [SessionAuthentication, JWTAuthentication]
   serializer_class = OrgSerializer


   def get_queryset(self):
       user_org = self.request.user.organization
       user_role = self.request.user.role.name
       role_based_queryset = {
           "Supervisor": Organization.objects.all(),
           "Reseller": Organization.objects.filter(
               Q(reseller=user_org) | Q(id=user_org.id)
           ),
           "ResellerAgent": Organization.objects.filter(agent=user_org),
           "CustomerAdmin": Organization.objects.filter(id=user_org.id),
       }
       return role_based_queryset.get(user_role)










Filters (install modules first)


filterset_fields = ["name"]
   search_fields = ["name"]



REST_FRAMEWORK = {
   "DEFAULT_AUTHENTICATION_CLASSES": [
       "rest_framework_simplejwt.authentication.JWTAuthentication",
       "rest_framework.authentication.SessionAuthentication",
   ],
   "DEFAULT_FILTER_BACKENDS": [
       "django_filters.rest_framework.DjangoFilterBackend",
       "rest_framework.filters.SearchFilter",
       "rest_framework.filters.OrderingFilter",
   ],
   "DEFAULT_PERMISSION_CLASSES": ("rest_framework.permissions.IsAuthenticated",),
   "DEFAULT_PAGINATION_CLASS": "rest_framework.pagination.PageNumberPagination",
   "PAGE_SIZE": 100,
   "DEFAULT_SCHEMA_CLASS": "rest_framework.schemas.coreapi.AutoSchema",
}





Send user info in token


class MyTokenObtainPairSerializer(TokenObtainPairSerializer):
   @classmethod
   def get_token(cls, user):
       token = super().get_token(user)
      
       token["username"] = user.email
       token["is_staff"] = user.is_staff
       token["organization"] = user.organization.id
       token["user_type"] = user.role.name
     


  
      
       return token


class MyTokenObtainPairView(TokenObtainPairView):
   serializer_class = MyTokenObtainPairSerializer









LIST Input from serializer

class PrintQRcodeSerilizer(serializers.Serializer):
   qrcodes = serializers.ListField(child=serializers.IntegerField(min_value=0, max_value=100))
   




class printqrcodeapiview(ModelViewSet):
   permission_classes = [WriteByAdminOnlyPermission, ReadByAdminOnlyPermission]
   serializer_class = PrintQRcodeSerilizer
   queryset = ""




   def create(self, request, *args, **kwargs):
       serializer = PrintQRcodeSerilizer(data= request.data)
       if serializer.is_valid(raise_exception=True):
           qty = serializer.validated_data["qrcodes"]
           print("-----------------",type(qty))  
           return "Done"


















Request (some mata data) - HTTP REQUEST 
 Apache/nginx
Wsgi (web server gateway interface -> how web server communicate  
               To application)
Middleware

Setting.py  -> ROOT_URLCONF  =”project.urls” (start urls) first come

urls.py

Urlpatterns list   scan top to bottom

Find view from imported file  - request is first parameter of view func

database (models.py ) perform some sql operations 

return a http response





conf.py

path function get two args 

def _path(route, view, kwargs=None, name=None, Pattern=None):



Resolver.py

Check url pattern
def _check_callback(self):










https://raturi.in/blog/custom-mixins-django-class-based-views/







@action
from rest_framework.decorators import action

@action(detail=True, methods=["get"], name="user_posts")
   def user_posts(self, request, pk):
       target_user = get_object_or_404(User, pk=pk)
       data = []
       if (
           target_user != request.user
           and target_user.is_private
           and not FollowRequest.is_friends(target_user, request.user)
       ):
           return Response({"data": data})
       queryset = self.queryset.filter(created_by=target_user)
       page = self.paginate_queryset(queryset)
       if page is not None:
           serializer = self.get_serializer(page, many=True)
           return self.get_paginated_response(serializer.data)


       serializer = self.get_serializer(queryset, many=True)
       return Response(serializer.data)




My MiddleWARE
def my_middle(get_response):
  print("one time")
  def myfunction(request):
     response = get_response
     print("After Response")
     return response
  return myfunction




"apps.organization.middle.my_middle"









Django core 
https://docs.djangoproject.com/en/4.1/ref/settings/#core-settings


STATIC
urlpatterns += static(settings.MEDIA_URL, document_root=settings.MEDIA_ROOT) + static(
   settings.STATIC_URL, document_root=settings.STATIC_ROOT
)




STATIC_URL = "/static/"
MEDIA_URL = "/media/"


MEDIA_ROOT = os.path.join(BASE_DIR, "media")


STATIC_ROOT = BASE_DIR / "staticfiles"




DJANGO_SETTINGS_MODULE = "RoboKidz.settings.production"


All values Serializer



from rest_framework import serializers


from .models import Song, Artist







class ArtistSerializer(serialisers.ModelSerializer):


 songs = serializers.PrimaryKeyRelatedField(


   many=True,


   read_only=False,


   queryset=Song.objects.all()


 )



SHOW IMG IN ADMIN

models.py


   def image_tag(self):
       from django.utils.html import escape, mark_safe
       try:


           return mark_safe(
               '<img src="%s" width="80px" height="80px"/>' % escape(self.qrcode.url)
           )
       except Exception as e:
           pass


   image_tag.short_description = "Image"
   image_tag.allow_tags = True




Admin.py




@admin.register(QRCode)
class QrCodeAdmin(admin.ModelAdmin):
   fields = ('category',)
   list_display = ("id", "qr_sr_no",  "category", "printdate", "image_tag")  #qrcode
   readonly_fields = ( "qr_sr_no", "image_tag")







DOWNLOAD BOTO3 
from django.test import TestCase
import os


# Create your tests here.
import os


import boto3
from boto3.session import Session


AWS_ACCESS_KEY_ID = ""
AWS_SECRET_ACCESS_KEY = ""




session = Session(aws_access_key_id= AWS_ACCESS_KEY_ID , aws_secret_access_key="")
s3 = session.resource('s3')
Service_Name = "S3"
bu = "clapzbucket"


mybucket = s3.Bucket(bu)




for i in  mybucket.objects.all():
   print(i)




USER MEDIA MODELVIEWSET URL

context={"request": self.context["request"]}


Access_KEY_ID = ""
Access_Key = "”
Service_Name = "S3"
Region_name= "ap-south-1"



Request user 
user = self.context["request"].user

Seilizer two
class CategorySerializer(serializers.ModelSerializer):
   videos = serializers.SerializerMethodField()


   def get_videos(self, obj):
       videos = obj.video_category.all()
       return VideoResponseserilizer(videos, many=True).data


   class Meta:
       model = Category
       fields = ("id", "name", "videos")





Data Paload 






class SubscriptionApiView(viewsets.ModelViewSet):
   queryset = Subscription.objects.all()
   serializer_class = SubscriptionSerializer


   def get_serializer(self, *args, **kwargs):
       if self.request.method in ["POST", "PATCH", "PUT"]:
           serializer_class = self.get_serializer_class()
       else:
           serializer_class = SubscriptionResponseSerializer


       kwargs.setdefault("context", self.get_serializer_context())
       return serializer_class(*args, **kwargs)


   def create_task_payload(self, before, after):
       output = []
       for i in before.keys():
           output.append({
               "name":i,
               "before":before[i],
               "after":after[i]
           })
       print(output)


       return output


   def update(self, request, *args, **kwargs):
       partial = kwargs.pop('partial', False)
       instance = self.get_object()
       befor_save_data= Subscription.objects.filter(id=int(self.kwargs["pk"])).first()
       befor_save_data = dict(SubscriptionSerializer(instance=befor_save_data).data)
       serializer =SubscriptionSerializer(instance, data=request.data, partial=partial)
       serializer.is_valid(raise_exception=True)
       self.perform_update(serializer)
       after_save_data = dict(serializer.data)
       print("------",type(befor_save_data))
       print()
       print("------",type(after_save_data))
      
       newdatapayload = self.create_task_payload(befor_save_data, after_save_data)
       # print(newdatapayload)
   


  
       return  Response(serializer.data)

My code 






class SubscriptionApiView(viewsets.ModelViewSet):
   queryset = Subscription.objects.all()
   serializer_class = SubscriptionSerializer


   def get_serializer(self, *args, **kwargs):
       if self.request.method in ["POST", "PATCH", "PUT"]:
           serializer_class = self.get_serializer_class()
       else:
           serializer_class = SubscriptionResponseSerializer


       kwargs.setdefault("context", self.get_serializer_context())
       return serializer_class(*args, **kwargs)


   def update(self, request, *args, **kwargs):
       partial = kwargs.pop('partial', False)
       instance = self.get_object()
       obj= Subscription.objects.filter(id=int(self.kwargs["pk"])).values()
       olddata = SubscriptionSerializer(obj,many=True)


       olddata = list(obj)
       serializer =SubscriptionSerializer(instance, data=request.data, partial=partial)
       serializer.is_valid(raise_exception=True)
       self.perform_update(serializer)
       obj= Subscription.objects.filter(id=int(self.kwargs["pk"])).values()
       newdata = list(obj.values())


       olddict = {}
       for i in olddata:
           for key,values in i.items():
               olddict[key] = values
       print(olddict)


       newdict = {}
       for j in newdata:
           for key,values in j.items():
               olddict[key]= values
       print(newdict)




       def data(olddict,newdict):
           kk =[]
           for i in olddict.keys():
               kk.append({
                   "name": i,
                   "before":olddict[i],
                   "after":newdict[i],
              


               })
              
           print(kk)
           return kk


       jj = data(olddict,newdict)



Data base

def Insert(request):
    kk = Students.objects.filter(FirstName="sam").values()
    lst = list(kk)
    ja = Students.objects.filter(FirstName="sam")


    kmm =  [{'id': 2,
            'FirstName': 'sam',
            'LastName': 'kumar',
            'Email': 'sam@gmail.om',
            'Contact': 12333}]




   


    for i in kmm:
        print(i)

 





OVERRIDE UPDATE method in modelviewset


class SubscriptionApiView(viewsets.ModelViewSet):
   queryset = Subscription.objects.all()
   serializer_class = SubscriptionSerializer


   def get_serializer(self, *args, **kwargs):
       if self.request.method in ["POST", "PATCH", "PUT"]:
           serializer_class = self.get_serializer_class()
       else:
           serializer_class = SubscriptionResponseSerializer


       kwargs.setdefault("context", self.get_serializer_context())
       return serializer_class(*args, **kwargs)


   def update(self, request, *args, **kwargs):
       partial = kwargs.pop('partial', False)
       instance = self.get_object()
       before_data= Subscription.objects.filter(id=int(self.kwargs["pk"])).values()
       print("This is before data",before_data)
      
       serializer =SubscriptionSerializer(instance, data=request.data, partial=partial)
       serializer.is_valid(raise_exception=True)
       self.perform_update(serializer)
       after = serializer.data
       print()
       print()
       print("This is Afetr Data", after)




       return  Response(serializer.data)


     










Payload

class VerifyEmail(generics.GenericAPIView):
   permission_classes = []
   serializer_class = PrentEmailSerializer
   def get(self, request):
       token = request.GET.get("token")
       payload = jwt.decode(token, settings.SECRET_KEY, algorithms=['HS256'])       
       user = User.objects.get(id=payload['user_id'])
       if not user.parent_email_verified:
           user.parent_email_verified = True
           user.save()
           return Response({"email": "succefull activated"}, status=status.HTTP_200_OK)
       else:
           return Response(status=status.HTTP_400_BAD_REQUEST)
   

Update 


q = get_object_or_404(MyModel,pk=some_value)
q.field1 = 'some value'
q.save()





Expire plan celery

@shared_task(bind =True)
def ExpirePlan(self):
   try:
       time = datetime.now()
       my_string1 = time.isoformat(' ', 'seconds')
       print(my_string1)
       expire_obj = Subscription.objects.filter(subscription_close_date = my_string1)
       print("------------------",expire_obj)
       for obj in expire_obj:
           # email = expire_obj.email
           # subject = 'Expire RoboKidz Portal '
           # message = f' Your Portal acress Expire Today  {time}'
           # email_from = settings.EMAIL_HOST
           obj.active = True
           obj.save() 
           # send_mail(subject, message, email_from, [email])
           print("Run Sucessfully")
   except Exception as e:
       print(e)






# @shared_task(bind= True)
# def send_mail_task(self):
  
#     # print("Mail sending.......")
#     # subject = 'welcome to Celery world'
#     # message = 'This is celery test'
#     # email_from = EMAIL_HOST_USER
#     # recipient_list = ['user1@yopmail.com' ]
#     # send_mail( subject, message, email_from, recipient_list )
#     # return "Mail has been sent........"










app.conf.beat_schedule = {
       'ExpirePlan': {
       'task': 'apps.mobile.tasks.ExpirePlan',
       'schedule': 1.0, #every 30 seconds it will be called
       #'args': (2,) you can pass arguments also if rquired
       }
}






app.autodiscover_tasks()
@app.task(bind=True)
def debug_task(self):
   print(f'Request: {self.request!r}')








DELETE From django_celery_beat_periodictask;








Modelviewset Pk 
 def retrieve(self, request, *args, **kwargs):
       param = kwargs
       pk = param['pk']
       return Subscription.objects.filter(id=pk)
       return


       # return super().retrieve(request, *args, **kwargs)




Call API Data 


def get_products(request):
  response = requests.get("https://fakestoreapi.com/products").json()
  for i in response:
   print(i['price'])
   Products.objects.create(title=i['title'],price =i['price'],category=i['category'],description=i['description'], image=i['image'])
  return render(request,'show.html',{'response':response})


def get_single(request,id):
   response = requests.get("https://fakestoreapi.com/products/1").json()
  
   return render(request,'one.html',{'response':response})







Export data as json file with all apps

python3 manage.py dumpdata users  -o  apps/users/fixtures/my.json
(  python3 manage.py dumpdata users  -o  	
)


Load json data in DB (when we already json file) 


python3 manage.py loaddata apps/users/fixtures/role.json




Format in vs code  –  shift +alt + F  ( window)
                      Shift + option+F    - mac
                       Shift + alt +I   - linux
















Load fix ture at migrate time 






from django.conf import settings
from django.core.management import call_command




@receiver(post_migrate)
def load_fixture_data(sender, **kwargs):
   try:
 
       if  not settings.FIXTURES_LOADED:
           call_command('loaddata', 'apps/users/fixtures/group.json')
           settings.FIXTURES_LOADED = True
   except Exception as e:
      
       print(e)








#django rest change password (install django restassword)
from django_rest_passwordreset.views import ResetPasswordConfirm, GenericViewSet


from django_rest_passwordreset.views import ResetPasswordToken



Serilizer.py



class ChangePasswordSerializer(PasswordTokenSerializer):
   confirm_password = serializers.CharField(required=True)


   def validate(self, attrs):
       if attrs["password"] != attrs["confirm_password"]:
           raise serializers.ValidationError(
               {"password": "Password fields didn't match."}
           )
       return attrs
  



views.py



class ResetPasswordConfirmView(ResetPasswordConfirm):
   serializer_class = ChangePasswordSerializer


   """
   An Api ViewSet which provides a method to reset a password based on a unique token
   """


   def post(self, request, reset_password_token, *args, **kwargs):
       serializer = self.serializer_class(data=request.data)


       reset_password_token = ResetPasswordToken.objects.filter(
           key=reset_password_token).first()


       def __init__(self):
           print("init")


       if reset_password_token:
           return super().post(request, *args, **kwargs)


       else:
           return Response({'message': "invaild Request or Token"})


  



Urls
  path(
       "password_reset/",
       include("django_rest_passwordreset.urls", namespace="password_reset"),
   ),
   path('confirm/<reset_password_token>/',views.ResetPasswordConfirmView.as_view(),name="confirm_passowrd"),
  

Email.py 

from django.shortcuts import render,HttpResponse
# Create your views here.
from rest_framework.decorators import api_view
from rest_framework.response import Response
from rest_framework_simplejwt.views import TokenObtainPairView
from rest_framework import generics
from apps.organization.models import *
from rest_framework.permissions import AllowAny, IsAuthenticated
from rest_framework.decorators import api_view, permission_classes, authentication_classes
from rest_framework.authentication import SessionAuthentication


from .serializer import *
from . models import *










# Create your views here.




from my_backend.settings.base import *
from django_rest_passwordreset.signals import reset_password_token_created
from django_rest_passwordreset.signals import reset_password_token_created
from django.core.mail import send_mail 
@receiver(reset_password_token_created)
def password_reset_token_created(sender, instance, reset_password_token, *args, **kwargs):
   # email_plaintext_message = "http://127.0.0.1:8000{}?token={}".format(reverse('confirm/<token>/:confirm_passowrd'), reset_password_token.key)
   #     # title:


   email_plaintext_message = f"http://127.0.0.1:8000/confirm/{reset_password_token.key}"
 
   subject ="Password Reset for {title}".format(title="Some website title")
       # message:
   message = email_plaintext_message
       # from:
   email_from = EMAIL_HOST
  
       # to:
   send_mail(subject, message, email_from, [reset_password_token.user.email])
   print("send")








  




















Django command fack data generate

Appname /management/commands

#dummy data
import random


from faker import Faker
fake = Faker()
fake.random.seed(4321)
from django.core.management.base import BaseCommand
class Command(BaseCommand):
   help = "Display Info "


   def handle(self, *args,**kwargs ):
       list = ["supervisor","reseller","resellerAgent","customerAdmin"]
       for i in list:
            UserRoles.objects.create(name=i )


       print("roles add successfully")


   def handle(self, *args,**kwargs ):
       for i in range(4):
           name=fake.name()+str(random.randint(100,1234))   
           password = "admin@123"
           email = fake.email()
           choice = ["1","2","3","4"]
           role = random.choice(choice)
           cust = Customer.objects.create(name=name)
           user = SeteraUser.objects.create(
               first_name=name,email=email,
               role_id=role,
               customer_id = cust.id)
          
           user.set_password(password)
           user.save()
       print("Data Inserted Successfully")



Fake data generate 


import random
from random import shuffle
from faker import Faker


fake = Faker()
fake.random.seed(4321) 
def insertdummydata(request):
   for i in range(10):
       name=fake.name()
       print("-",name)
       password = "sajal@123"
       email = fake.email()
       choice = ["1","2","3","4"]
       role = random.choice(choice)
       cust = Customer.objects.create(name=name)
       user = SeteraUser.objects.create(
           first_name=name,email=email,
           role_id=role,
           customer_id = cust.id
       )
       user.set_password(password)
       user.save()
   return  HttpResponse("Done")
‘



Validate modelviewset(null)




   def validate_role(self, value):
   
       if not value:
           raise serializers.ValidationError(
               {"role": "Please Select the User Role"})   


       return value
     
          



System celery 

sudo systemctl start redis-server.service


Perform create



class QueryApiView(viewsets.ModelViewSet):
   queryset = QueryifAny.objects.all()
   serializer_class = QuerySerializer


   def perform_create(self, serializer):
       serializer.save(user=self.request.user, )







My profile



class MyProfileAPIView(generics.GenericAPIView):
   serializer_class = MyProfileSerializer


   def get(self, request):
       serializer = self.serializer_class(request.user, context={"request": request})
       return Response({"data": serializer.data})




Update otp 


class MobileOtp(generics.CreateAPIView):
   serializer_class = PhoneModelSerilalizer
   permission_classes = ()


   def post(self, request):
       serializer = self.serializer_class(data=request.data)
       if serializer.is_valid(raise_exception=True):
           mo =serializer.validated_data['mobile']
           if PhoneModel.objects.filter(mobile=mo).exists():
               instance = PhoneModel.objects.get(mobile=mo)
               instance.otp = random.randint(1000,9999)
               instance.save()
               content = {"mobile": instance.mobile, "otp": instance.otp}
               return Response(content, status=status.HTTP_201_CREATED)
           else:
               instance = serializer.save()
               content = {"mobile": instance.mobile, "otp": instance.otp}
               # send_otp(mobile,otp)
               return Response(content, status=status.HTTP_201_CREATED)
       return Response(
           {"message", "invalid number"}, status=status.HTTP_400_BAD_REQUEST
       )




Scan wp


root@jarvis1:/home/shubpy# wpscan --url https://www.shubpy.com / -e ap,at,cb ,dbe -o myscan.txt  --random-user-agent




Reset Password 




from django.db import models


# Create your models here.


from django.dispatch import receiver
from django.urls import reverse
from django_rest_passwordreset.signals import reset_password_token_created
from django.core.mail import send_mail 




@receiver(reset_password_token_created)
def password_reset_token_created(sender, instance, reset_password_token, *args, **kwargs):


   email_plaintext_message = "{}?token={}".format(reverse('password_reset:reset-password-request'), reset_password_token.key)


   send_mail(
       # title:
       "Password Reset for {title}".format(title="Some website title"),
       # message:
       email_plaintext_message,
       # from:
       "noreply@somehost.local",
       # to:
       [reset_password_token.user.email]
   )




736674 Validation Modelviewset

   def validate(self, attrs):
       qty = attrs.get("quantity")


       if qty <= 0 :
           raise serializers.ValidationError("Please Choose a valid number between 0 to 100")


       elif qty >= 100:
           raise serializers.ValidationError("You Can print 100 Qr only at a time")


       return super().validate(attrs)





Request.user in seriliazer 

User = self.context["request"].user


Random Password generate

import random


# Random string of length 5
result_str = ''.join((random.choice('abcdxyzpqr1234567890') for i in range(8)))
print(result_str)







Manny to many fields





Modrls.py


def get_videoclass(self):
       return ",".join([str(p) for p in self.videoclass.all()])








Admin.py
 def get_hashtag(self,obj):
       return [i.hashtag for i in obj.hashtag.all()]







class StudentDetailsAdmin(admin.ModelAdmin):
   list_display = ['first_name',
                       'is_verify',]
   list_editable = ['is_verify']


   def get_queryset(self, request):
       vuser = StudentDetails.objects.filter(is_verify= False)
       return vuser






@admin.register(PostHashTags)
class PostTagAdmin(admin.ModelAdmin):
   list_display = ('hashtag',)




@admin.register(Mypost)
class MyPostAdmin(admin.ModelAdmin):
   list_display = ('name',
                       'discription',
                       'get_hashtag',
                       )
   def get_hashtag(self,obj):
       return [i.hashtag for i in obj.hashtag.all()]



class ProfileAPIView(generics.ListAPIView):
   serializer_class = AllProfileSerializer


   def get_blocked_user_ids(self, data):
       ids = []
       for i in data:
           ids.append(i["created_by"])
           ids.append(i["blocked_user"])
     
       return set(ids)

BAsic List 
document = Document()




li = [3,5,6,7,8,9]
for i in li[3:5]:
   print("-------------",i)



My doc 


import qrcode
import random
import os, math
from docx.shared import Inches
import os




# import os
# file_path = "/home/shubpy/CirtificATE/new1111111111.png"
# os.startfile(file_path,'print')


# import subprocess, sys


# opener = "open" if sys.platform == "darwin" else "xdg-open"
# subprocess.call([opener, file_path])
from docx import Document
document = Document()










# document.save("/home/shubpy/QRCODE/hello.docx")
cat = input("Enter the category : ")
code = int(input("Enter the  range : "))
name=(cat[:3]).upper()
print(name)
res = []
for i in range(code):
   no = random.randint(10000,100000)
   imgname =f"RR{name}{no}"
   print(imgname)
   print(type(imgname))


   img = qrcode.make(data=imgname)
  
   path = img.save(f"/home/shubpy/QRCODE/qr/{imgname}.png")
   pa = f"/home/shubpy/QRCODE/qr/"
   for path in os.listdir(pa):
   # check if current path is a file
       if os.path.isfile(os.path.join(pa, path)):
           res.append(path)














document.add_heading('Grid Images', 0)


# add paragraph and get run
p = document.add_paragraph()
r = p.add_run()
# j = int(input("How many print QR code from database :"))
for i in res:
   print("------------------",i)
   r.add_picture(f"/home/shubpy/QRCODE/qr/{i}", width=Inches(1.46), height=Inches(1.45))
   r.add_text(" ")


document.save("/home/shubpy/QRCODE/doc/hello.docx")
      
  

2
Doc refer 



from docx.shared import Inches


# add image with height and width in Inches
doc.add_picture("images/emma-dau.jpg", width=Inches(5), height=Inches(3))






from docx.enum.text import WD_ALIGN_PARAGRAPH
# add picture
doc.add_picture("images/emma-dau.jpg", width=Inches(5), height=Inches(3))
# Get image paragraph and align
last_paragraph = doc.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER




# grid code here
from docx import Document
from docx.shared import Inches
import os
# create document
document = Document()
# add heading
document.add_heading('Grid Images', 0)
# add paragraph and get run
p = document.add_paragraph()
r = p.add_run()
# iterate over each image in directory
for image_name in os.listdir("images"):
   # add image so it creates a row with 4 images
   r.add_picture(f"images/{image_name}", width=Inches(1.46), height=Inches(1.45))
   r.add_text(" ") # add space for image seperation
# save document
document.save('demo.docx')












from docx.shared import Inches, Cm
import os
from docx.enum.text import WD_ALIGN_PARAGRAPH


doc = Document() # create doc
doc.add_heading('Images in Table', 0) # add heading


# create table with two rows and columns (Per row images are 3)
table = doc.add_table(rows=0, cols=3, style="Table Grid")
image_dir = "images"
images = os.listdir(image_dir )


for i in range(4): # show 4 rows of image (4x3 = 12 images)
   image_row = table.add_row() # add row to table for images
   cap_row = table.add_row() # add row to table for image name
   for j in range(3): # Iterate and get 3 images for row
       image_name = images.pop()


       # add image to table
       set_cell_margins(image_row.cells[j], top=100, start=100, bottom=100, end=50)
       # add image to cell and align center
       paragraph = image_row.cells[j].paragraphs[0]
       paragraph.add_run().add_picture(f"{image_dir }/{image_name}", width=Inches(1.8), height=Inches(1.6))
       paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


       # add caption to table
       cap_row.cells[j].text = image_name
       cap_row.cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER




# save to file
doc.save("images-table.docx")
  














File size and read Basic 



from django.test import TestCase


import os,math


# folder path
dir_path = r'/home/shubpy/Videos/demo/'


# list to store files
res = []


# Iterate directory
for path in os.listdir(dir_path):
   # check if current path is a file
   if os.path.isfile(os.path.join(dir_path, path)):
       res.append(path)


for i in res:
   file_stats = os.stat(f"/home/shubpy/Videos/demo/{i}")
   s =file_stats.st_size


   size_name = ("B", "KB", "MB", "GB", "TB", "PB", "EB", "ZB", "YB")
   i = int(math.floor(math.log(s, 1024)))
   p = math.pow(1024, i)
   s = round(s/ p, 2)
   print( f"%s %s" % (s, size_name[i]))









from django.test import TestCase


import os,math


# folder path
dir_path = r'/home/shubpy/Videos/demo/'


# list to store files
res = []


# Iterate directory
for path in os.listdir(dir_path):
   # check if current path is a file
   if os.path.isfile(os.path.join(dir_path, path)):
       res.append(path)


for i in res:
   # file_stats = os.stat(f"/home/shubpy/Videos/demo/{i}")
   # s =file_stats.st_size


   size_name = ("B", "KB", "MB", "GB", "TB", "PB", "EB", "ZB", "YB")
   i = int(math.floor(math.log(1e+12, 1024)))
   p = math.pow(1024, i)
   s = round(1e+12/ p, 2)
   print( f"%s %s" % (s, size_name[i]))







File size (st_size=8615175)

from django.test import TestCase


import os


# folder path
dir_path = r'/home/shubpy/Videos/demo/'


# list to store files
res = []


# Iterate directory
for path in os.listdir(dir_path):
   # check if current path is a file
   if os.path.isfile(os.path.join(dir_path, path)):
       res.append(path)


for i in res:
   file_stats = os.stat(f"/home/shubpy/Videos/demo/{i}")
   print(file_stats)







User read files from directoers


from django.test import TestCase


import os


# folder path
dir_path = r'/home/shubpy/Videos/demo/'


# list to store files
res = []


# Iterate directory
for path in os.listdir(dir_path):
   # check if current path is a file
   if os.path.isfile(os.path.join(dir_path, path)):
       res.append(path)
print(res)







Read QRcode Image
import glob
import cv2
import pandas as pd
import pathlib
print("-----------------")




def read_qr_code(filename):
   """Read an image and read the QR code.
  
   Args:
       filename (string): Path to file
  
   Returns:
       qr (string): Value from QR code
   """
  
   try:
       img = cv2.imread(filename)
       detect = cv2.QRCodeDetector()
       value, points, straight_qrcode = detect.detectAndDecode(img)
       return value
   except:
       return




value = read_qr_code('/home/shubpy/projects/SOW/ROBO-SERVER/RoboKidz/RoboKidz/media/images/RRHAN29570.jpg')


print(value)




QR CODE BASIC 

from django.test import TestCase


import os


# folder path
dir_path = r'/home/shubpy/Videos/demo/'


# list to store files
res = []


# Iterate directory
for path in os.listdir(dir_path):
   # check if current path is a file
   if os.path.isfile(os.path.join(dir_path, path)):
       res.append(path)
print(res)






Main —---------------------------------




import qrcode
import random
cat = input("Enter the category : ")
code = int(input("Enter the  range : "))
name=(cat[:3]).upper()
print(name)


for i in range(code):
   no = random.randint(10000,100000)
   imgname =f"RR{name}{no}"
   print(imgname)
   print(type(imgname))


   img = qrcode.make(data=imgname)
   img.save(f"{imgname}.png")
  








































from enum import Enum
import random  






cat = input("Enetr the catagory name : ")
category  =(cat[:3]).upper()


code = int(input("Eneter the  num :"))




title="Robokidz API"
print("this is tpe :  ",type(title))








import qrcode 
obj_qr = qrcode.QRCode( 
   version = 1, 
   error_correction = qrcode.constants.ERROR_CORRECT_L, 
   box_size = 10, 
   border = 4, 
) 


# using the add_data() function 


for i in range(code):
   no =random.randint(10000,100000)
   obj_qr.add_data(res) 
   # using the make() function 
   obj_qr.make(fit = True) 
   # using the make_image() function 
   qr_img = obj_qr.make_image(fill_color = "cyan", back_color = "black") 
   # saving the QR code image 


   qr_img.save(f"RR{category}{no}.png") 












Queyset values 
   def get_queryset(self):
       h = Post.objects.filter(created_by=self.request.user)


       j = [i.is_approved  for i in h]
       print("This is jj values ---",j)









Swaggers Setup (run migrations is madentory)


pip install django-rest-swagger






pip install django-rest-swagger


pip install drf-yash


from django.urls import path,include,re_path






from rest_framework_swagger.views import get_swagger_view
schema_view = get_swagger_view(title='Rest API Sajal ViEW')


from rest_framework import permissions
from drf_yasg.views import get_schema_view
from drf_yasg import openapi




schema_view = get_schema_view(
   openapi.Info(
       title="Robokidz API",
       default_version="v1",
       description="Robokidz api documentation LCM",
   ),
   public=True,
   permission_classes=(permissions.AllowAny,),
)








urlpatterns = [


   re_path(
       r"^doc(?P<format>\.json|\.yaml)$",
       schema_view.without_ui(cache_timeout=0),
       name="schema-json",
   ),  # <-- Here
   path(
       "doc/",
       schema_view.with_ui("swagger", cache_timeout=0),
       name="schema-swagger-ui",
   ),  # <-- Here
   path(
       "redoc/", schema_view.with_ui("redoc", cache_timeout=0), name="schema-redoc"
   ),






   path('admin/', admin.site.urls),
   path('',include('apps.user.urls'))


]



There is one error occur if 









Refer code 


Retun 2 serilizrs data 





class SearchAPIView(generics.GenericAPIView):
   permission_classes = (IsAuthenticated,)
   serializer_class = SearchSerializer


   def post(self, request, *args, **kwargs):
       serializer = self.serializer_class(data=request.data)
       serializer.is_valid()
       # search_val = serializer.validated_data["search"]
       search_val = request.data["search"]
       posts = Post.objects.filter(
           Q(caption__icontains=search_val,is_approved=True )
           | Q(description__icontains=search_val, is_approved=True)
           | Q(created_by__mobile__icontains=search_val, is_approved=True)
           | Q(created_by__username__icontains=search_val,is_approved=True)
           | Q(created_by__name__icontains=search_val, is_approved=True)
           | Q(created_by__id__contains=search_val,is_approved=True)
           | Q(created_by__email__icontains=search_val,is_approved=True)
           | Q(created_by__first_name__icontains=search_val,is_approved=True)
       )


       users = User.objects.filter(
           Q(user_type__icontains=search_val)
           | Q(username__icontains=search_val)
           | Q(email__icontains=search_val)
           | Q(school__name=search_val)
           | Q(first_name__icontains=search_val)
           | Q(name__icontains=search_val)
           # | Q(mobile=search_val)
           | Q(state__icontains=search_val)
           | Q(city__icontains=search_val)
           | Q(grade__name=search_val)
           | Q(parent_email=search_val)
           | Q(address__icontains=search_val)
           | Q(workspace__icontains=search_val)
           | Q(position__icontains=search_val)
           | Q(qualification__icontains=search_val)
       )


      
       posts_data = [PostResponseSerializer(posts, context={"request": request}, many=True).data,  AllProfileSerializer(users, context={"request": request}, many=True).data]


       return Response({"posts": posts_data})











class CrateCirtificates(viewsets.ModelViewSet):
   queryset=Certificate.objects.all()
   serializer_class = CreateCertificateSerializer






class GenrateUserCirtificate(viewsets.ModelViewSet):
   serializer_class=GenerateUserCertificateSerilizer
   queryset=PostCertificate.objects.all()

















from PIL import Image, ImageDraw, ImageFont
import datetime
import uuid


class GenerateUserCertificateSerilizer(serializers.ModelSerializer):
   class Meta:
       model = PostCertificate
       fields = ["post", "user", "follow", "certificate", "created_by", "updated_by"]
       print(fields)


   def create(self, validated_data):
       post = validated_data["post"].id
       user = validated_data["user"].id
       follow = validated_data["follow"].id
       certificate = validated_data["certificate"].id
       certificate = validated_data["certificate"].certificate
       img = Image.open(certificate)
       now = datetime.datetime.now()
       today = now.strftime("%d-%m-%Y")


       # conditions here


       image_editable = ImageDraw.Draw(img)


       font = ImageFont.truetype("E:/PythonPillow/Fonts/FreeMono.ttf", 48)
       image_editable.text(
           (500, 500), validated_data["user"].first_name, fill=(0, 0, 0), font=font
       )
       image_editable.text( (530, 700), "Outstanding followers ", fill=(0, 0, 0), font=font )


       image_editable.text((950, 800), "Follow", fill=(0, 0, 0), font=font)


       font = ImageFont.truetype("E:/PythonPillow/Fonts/FreeMono.ttf", 38)
       image_editable.text((500, 900), today, fill=(0, 0, 0), font=font)
       img.show()
       file_name = uuid.uuid4().hex
       img.save(f"RoboKidz/media/usercirtificate/Cirtificate{file_name}.png")
       validated_data["certificate_image"] = f"/usercirtificate/Cirtificate{file_name}.png"
       RA = PostCertificate.objects.create(**validated_data)
       return RA

























Validate foreign key values 



class GenerateUserCertificateSerilizer(serializers.ModelSerializer):
   class Meta:
       model = PostCertificate
       fields = ["post", "user", "follow", "certificate"]
       print(fields)


   def create(self, validated_data):
       Image = validated_data["certificate"].certificate
       print("----------------------",Image)
      


       img_info = PostCertificate.objects.create(**validated_data)
       return img_info





























Update python3 

https://www.debugpoint.com/install-python-3-11-ubuntu/


sudo apt install software-properties-common

sudo add-apt-repository ppa:deadsnakes/ppa
sudo apt update 

sudo apt install python3.11
sudo update-alternatives --install /usr/bin/python3 python3 /usr/bin/python3.11 2


sudo update-alternatives --config python3


curl -sS https://bootstrap.pypa.io/get-pip.py | python3.11




https://www.debugpoint.com/install-python-3-11-ubuntu/






class EquipmentType(models.Model):
   equipment_type = models.CharField(verbose_name="Equipment Type", max_length=50, unique=True)


   def __unicode__(self):
       return self.equipment_type




class EquipmentManufacturer(models.Model):


   manufacturer_name = models.CharField(verbose_name="Manufacturer Name", max_length=50, unique=True)


   def __unicode__(self):
       return self.manufacturer_name




class EquipmentInfo(models.Model):


   equipment_type = models.ForeignKey(EquipmentType, verbose_name="Equipment Type")
   part_identifier = models.CharField(verbose_name="Machine ID (alias)", max_length=25)
   manufacturer_name = models.ForeignKey(EquipmentManufacturer, verbose_name="Manufacturer Name")
   serial_number = models.CharField(verbose_name="Serial Number", max_length=25)
   date_of_manufacture = models.DateField(verbose_name="Date of Manufacture", default=date.today)
   is_active = models.BooleanField(verbose_name="Is Active", default=True)


   def __unicode__(self):
       return self.part_identifier
serializers.py


class EquipmentTypeSerializer(serializers.ModelSerializer):
   class Meta:
       model = EquipmentType
       fields = ('id', 'equipment_type',)


class EquipmentManufacturerSerializer(serializers.ModelSerializer):
   class Meta:
       model = EquipmentManufacturer
       fields = ('id', 'manufacturer_name',)


class EquipmentInfoSerializer(serializers.ModelSerializer):
   class Meta:
       model = EquipmentInfo
       fields = ('id', 'equipment_type', 'part_identifier', 'manufacturer_name','serial_number', 'date_of_manufacture', 'is_active')


   equipment_type = EquipmentTypeSerializer(many=False)
   manufacturer_name = EquipmentManufacturerSerializer(many=False)


   def create(self, validated_data):
       equipment_type = validated_data.pop('equipment_type')
       manufacturer_name = validated_data.pop('manufacturer_name')
       equipment_info = EquipmentInfo.objects.create(**validated_data)
       return equipment_info
Assuming I already have relevant EquipmentType and EquipmentManufacturer objects created, I would like to add another EquipmentInfo object. What is the appropriate way to set up my EquipmentInfo serializer so that I can pass in information such as


{
"equipment_type":{
 "equipment_type":"already_created",
},
"part_identifier":"something_new",
"manufacturer_name":{
 "manufacturer_name":"already_created"
},
"serial_number":"WBA1",
"date_of_manufacture": "1900-01-01",
"is_active":true
}
or even better:


{
"equipment_type":"already_created",
"part_identifier":"something_new",
"manufacturer_name":"already_created",
"serial_number":"WBA1",
"date_of_manufacture": "1900-01-01",
"is_active":true
}




# @api_view(['GET'])
# @permission_classes([IsAuthenticated])
# @authentication_classes([SessionAuthentication, JWTAuthentication])
# def getRoutes(request):
#     routes = [
#         '/api/token/',
#         '/api/register/',
#         '/api/token/refresh/',
#         '/api/organization/'
#     ]
#     return Response(routes)


Urls
#     path('', views.getRoutes),



NOTIFICATION 




class notification(viewsets.ModelViewSet):
   serializer_class = NotificationSerializer
   queryset = Notification.objects.all()
   http_method_names: List[str] = ["get", "patch"]


   def get_serializer(self, *args, **kwargs):
       if self.request.method in ["POST","PATCH"]:
           serializer_class = self.get_serializer_class()
       else:
           serializer_class = NotificationResponseSerializer
       kwargs.setdefault("context", self.get_serializer_context())
       return serializer_class(*args, **kwargs)


   def get_queryset(self):
       return Notification.objects.filter(notified_to=self.request.user.id)
       





Noti signal




from django.db.models.signals import post_save, pre_save


# from django.contrib.auth.models import User
from django.dispatch import receiver
from apps.post.models import Like, LikeComment, Post, Comment
from .models import Notification
from apps.user.models import FollowRequest
from django.db.models import signals




@receiver(post_save, sender=Like)
def post_save_like(sender, instance, created, **kwargs):
   if instance.created_by==instance.post.created_by:
       return False
   Notification.objects.create(
       performed_by=instance.created_by,
       notified_to=instance.post.created_by,
       message=instance.created_by.username + " " +" Liked on your posts.",
       message_key="post like",
   )




@receiver(post_save, sender=LikeComment)
def post_save_like_comment(sender, instance, created, **kwargs):
   if instance.created_by==instance.comment.created_by:
       return False
   Notification.objects.create(
       performed_by=instance.created_by,
       notified_to=instance.comment.created_by,
       message=instance.created_by.username + " " +" Liked  on one  your comment. ",
       message_key="comment like",
   )




@receiver(post_save, sender=FollowRequest)
def post_save_follow_request(sender, instance, created, **kwargs):
   if instance.status == "pending":
       Notification.objects.create(
           performed_by=instance.created_by,
           notified_to=instance.receiver,
           message=instance.created_by.username + " " +"send your follow request",
           message_key="follow-request-sent",
       )




@receiver(post_save, sender=FollowRequest)
def post_save_accept_follow_request(sender, instance, created, **kwargs):
   if instance.status == "accepted":
       Notification.objects.create(
           performed_by=instance.created_by,
           notified_to=instance.receiver,
           message=instance.created_by.username + " " + "accepted your follow request",
           message_key="follow-request-accepted",
       )




@receiver(post_save, sender=Comment)
def post_save_comment(sender, instance, created, **kwargs):
   if instance.created_by==instance.post.created_by:
       return False
   Notification.objects.create(
       performed_by=instance.created_by,
       notified_to=instance.post.created_by,
       message=instance.created_by.username +" "+ "commented on  your post. ",
       message_key="comment send",
   )







MOdelview set Crud 


from django.shortcuts import render
from rest_framework import viewsets,status
from rest_framework.response import Response
from rest_framework.authtoken.models import Token
from django.contrib.auth.models import User
from django.db import transaction
from adminapp.serializers import UserSerializer
from myapp.authentication.authentication import cTokenAuthentication
from rest_framework.permissions import DjangoModelPermissions
from django.utils import timezone
# Create your views here.


class adminLogin(viewsets.ModelViewSet):
    queryset = User.objects.all()
    http_method_names = ['post']


    def create(self, request, *args, **kwargs):
        data = request.data
        try:
            with transaction.atomic():
                userData = User.objects.get(
                    username =   data["username"],
                )
                if userData.is_superuser == 0:
                    return Response({"message":"Only admin can login", "status": False,
                             "response": "fail", }, status=status.HTTP_400_BAD_REQUEST)


                if userData.check_password(data["password"]):
                    if Token.objects.filter(user=userData).exists():
                        Token.objects.get(user=userData).delete()


                    userData.last_login = timezone.now()
                    userData.save()
                    token=Token.objects.create(user=userData)
               
                return Response({"username":data['username'], "message": "Login Successfully.Please use token for further process",
                                "token":token.key,
                                    "status": True, "response": "success", }, status=status.HTTP_200_OK)
        except Exception as error:
            return Response({"message":str(error), "status": False,
                             "response": "fail", }, status=status.HTTP_400_BAD_REQUEST)


class UserViewSet(viewsets.ModelViewSet):
    authentication_classes = [cTokenAuthentication]
    permission_classes = [DjangoModelPermissions]
    queryset = User.objects.all()
    serializer_class=UserSerializer
    http_method_names = ['post','get','put']


    def create(self, request, *args, **kwargs):
        data = request.data
        try:
            with transaction.atomic():
               


                userData = User.objects.create(
                    username =   data["username"],
                    email = data["email"],
                    first_name =  data["first_name"],
                    last_name =  data["last_name"],
                    is_active = True,
                )
               
                userData.set_password(data["password"])
                userData.save()
               
                return Response({"username":data['username'], "message": "Your registration has been successfully completed.",
                                    "status": True, "response": "success", }, status=status.HTTP_200_OK)
        except Exception as error:
             return Response({"message":str(error), "status": False,
                             "response": "fail", }, status=status.HTTP_400_BAD_REQUEST)


    def update(self, request, *args, **kwargs):
        data = request.data
        try:
            with transaction.atomic():
                userData = User.objects.filter(id=int(self.kwargs['pk'])).update(
                    email = data["email"],


                    first_name =  data["first_name"],
                    last_name =  data["last_name"],
                )


                user=User.objects.get(id=int(self.kwargs['pk']))
                               
                user.set_password(data["password"])
                user.save()
               
                return Response({"message": "User updated successfully",
                                    "status": True, "response": "success", }, status=status.HTTP_200_OK)
        except Exception as error:
             return Response({"message":str(error), "status": False,
                             "response": "fail", }, status=status.HTTP_400_BAD_REQUEST)
   
    def list(self, request, *args, **kwargs):
        try:
            with transaction.atomic():
                userData = User.objects.all()


                data=self.serializer_class(userData,many=True).data
                return Response({"data":data,"message": "list",
                                    "status": True, "response": "success", }, status=status.HTTP_200_OK)
        except Exception as error:
             return Response({"message":str(error), "status": False,
                             "response": "fail", }, status=status.HTTP_400_BAD_REQUEST)






Two serializer view Response 


class  SubscriptionApiView(viewsets.ModelViewSet):
   queryset = Subscription.objects.all()
   serializer_class = SubscriptionSerializer
   def get_serializer(self, *args, **kwargs):
       if self.request.method in ["POST", "PATCH"]:
           serializer_class = self.get_serializer_class()
       else:
           serializer_class = SubscriptionResponseSerializer


       kwargs.setdefault("context", self.get_serializer_context())
       return serializer_class(*args, **kwargs)








# Signals

from django.shortcuts import render,HttpResponse
from . models import *
# Create your views here.
from rest_framework import generics
from rest_framework.response import Response


from . serializer import *


def Insert(request):
    kk = Students.objects.filter(FirstName="sam").values()
    lst = list(kk)
    ja = Students.objects.filter(FirstName="sam")




    kmm =  [{'id': 2,
            'FirstName': 'sam',
            'LastName': 'kumar',
            'Email': 'sam@gmail.om',
            'Contact': 12333}]


    ll = {}
    for i in kmm:
        for key,values in i.items():
            ll["before " +key] = values


    print(ll)
           


           


# Driver code




   
    # jam = list(kk)
    # print("=====================",jam)


    return render(request,"app\index.html")




class Mystudents(generics.ListCreateAPIView):
    queryset = Students.objects.all()
    serializer_class = StudentSerilizer


    def create(self,request,*args,**kwargs):
        serializer = StudentSerilizer(data =request.data)
        if serializer.is_valid():
            kk = serializer.save()
            # jam =Students.objects.filter(id=1)
            # print("------------------------------",jam)
           
            jj = serializer.data
            print(jj)
            print("-------------",type(jj))


            for key,values in jj.items():
                print(key,values)


            return Response(serializer.data )


        else:
            return Response({'msg':'mmm'})






   








Forget  passwsword custom 




def Change_password(email,token):
   try:
       subject = 'Reset Password'
       message = f"http://127.0.0.1:8000/confirm/{token}"
       email_from = EMAIL_HOST
       send_mail(subject, message, email_from, [email])
       print("send ----------------------------")
   except Exception as e:
       print("not send ",e)








import uuid
class ForgetPasswordSerializer(serializers.Serializer):
   email = serializers.EmailField()
   def validate(self, attrs):
       email = attrs.get('email')
       # print(email)
       obj = User.objects.filter(email = email).exists()
       if not obj:
           raise ValidationError("Email Not Exists")
       return super().validate(attrs)
  
  
  
  
class VerifyPasswordSerializer(serializers.Serializer):
   password = serializers.CharField(label=_("Password"), style={'input_type': 'password'})
   confirm_password = serializers.CharField(label=_("confirm_password"), style={'input_type': 'password'})
   token  = serializers.CharField()


   def validate_password(self, value):
       data = self.get_initial()
       password = data.get('confirm_password')
       confirm_password = value
       if password != confirm_password:
           raise ValidationError('Passwords and Confirm password not  matching')
       return value







class ForgetApiView(generics.CreateAPIView):
   serializer_class = ForgetPasswordSerializer
   permission_classes = ()


   def create(self, request, *args, **kwargs):
       token =  uuid.uuid1()
       serializer =  self.serializer_class(data=request.data)
       if serializer.is_valid(raise_exception=True):
           email = serializer.data['email']
           print("This is email , ", email)
           reset_user = User.objects.filter(email =email).first()
           if EmailToken.objects.filter(user = reset_user.id ).exists():
               instance = EmailToken.objects.get()
               instance.email_token = token
               token  = instance.email_token
               email = instance.user.email
               Change_password(email,token)
               instance.save()
               Change_password(instance.user.email,instance.email_token )
           else:
               obj = EmailToken.objects.create(email_token = token,user_id = reset_user.id)
               token  = obj.email_token
               email = obj.user.email
               Change_password(email,token)


           return Response({"message", "Email Send Succefully"} )
          
       return Response(
           {"message", "Not valid Email"}, status=status.HTTP_400_BAD_REQUEST )


class VerifyResetPassword(generics.CreateAPIView):
   queryset = ""
   serializer_class = VerifyPasswordSerializer
   permission_classes = ()


   def post(self, request,token, *args, **kwargs):
       serializer =  VerifyPasswordSerializer(data=request.data)
       if serializer.is_valid(raise_exception=True):
           token = serializer.data['token']
           newpassword = serializer.data['password']
           print()
           print(token)
           obj_pass = EmailToken.objects.filter(email_token = token).first()
           # print(obj_pass)
           if obj_pass:
                   updatepassword = User.objects.filter(id = obj_pass.user.id).first()
                   updatepassword.set_password(newpassword)
                   updatepassword.save()
                   obj_pass.delete()
                   return Response({"message", "Password Update Successfully"}, status=status.HTTP_202_ACCEPTED)


       return Response({"message", "Not valid Email"}, status=status.HTTP_400_BAD_REQUEST )
 



class MobileChangePassword(generics.CreateAPIView):
   queryset = ""
   serializer_class = MobileChangePasswordSerializer
   permission_classes = ()


   def post(self, request, *args, **kwargs):
       serializer =  MobileChangePasswordSerializer(data=request.data)
       if serializer.is_valid(raise_exception=True):
           otp = serializer.data['otp']
           newpassword = serializer.data['password']
           obj_pass = PhoneModel.objects.filter(otp = otp).first()
           if obj_pass:
               updatepassword = User.objects.filter(mobile = obj_pass.mobile).first()
               updatepassword.set_password(newpassword)
               updatepassword.save()
               obj_pass.delete()
               p = obj_pass.mobile
           return Response({"message", "Password Update Successfully"}, status=status.HTTP_202_ACCEPTED)


       
       return Response({"message", "Not valid Email"}, status=status.HTTP_400_BAD_REQUEST )






CORS HEADER


CORS_ORIGIN_ALLOW_ALL = True


CORS_ALLOWED_ORIGINS = [
   "http://localhost:3000",
   "http://localhost:8000",
   "http://127.0.0.1:3000",
   "http://127.0.0.1:8000",
]


CORS_ORIGIN_WHITELIST = (
   "http://localhost:8000",
   "http://127.0.0.1:8000",
   "http://localhost:3000",
   "http://127.0.0.1:3000",
)


CORS_ALLOW_HEADERS = [
   "access-control-allow-origin",
   "content-type",
   "authorization",
]






CORS_ALLOW_HEADERS = [
   'accept',
   'accept-encoding',
   'authorization',
   'content-type',
   'dnt',
   'origin',
   'user-agent',
   'x-csrftoken',
   'x-requested-with',
   'Access-Control-Allow-Origin',
]




